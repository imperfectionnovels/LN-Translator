"""Initiative 7 — DOCX import via /api/translate/upload.

Builds minimal DOCX fixtures in-memory via python-docx, posts them through
/upload, and asserts chapter splitting follows the same conventions TXT
imports use (heading lines like "Chapter N: …" or 第N章 become chapter
markers; the body parser carries unstyled-heading variants forward).
"""

import os
import sqlite3
from io import BytesIO
from pathlib import Path

import pytest
from fastapi.testclient import TestClient

from backend.db import SCHEMA
from backend.main import app

DB_PATH = Path(os.environ["DB_PATH"])


@pytest.fixture
def client(monkeypatch, tmp_path):
    if DB_PATH.exists():
        DB_PATH.unlink()
    conn = sqlite3.connect(DB_PATH)
    conn.executescript(SCHEMA)
    conn.commit()
    conn.close()

    monkeypatch.setattr("backend.services.covers.USER_DATA_ROOT", tmp_path)
    monkeypatch.setattr(
        "backend.services.covers._COVERS_DIR", tmp_path / "covers"
    )

    async def _noop_run(novel_id, chapter_id):
        return

    monkeypatch.setattr("backend.services.queue._run_translate", _noop_run)
    return TestClient(app)


def _build_docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    """Build a DOCX with `(style_name, text)` paragraphs. style_name=''
    means default paragraph style."""
    from docx import Document  # python-docx

    doc = Document()
    for style, text in paragraphs:
        p = doc.add_paragraph(text)
        if style:
            try:
                p.style = doc.styles[style]
            except KeyError:
                # Style not present in the default template — fall back to
                # adding a heading directly.
                if style.startswith("Heading"):
                    level = int(style.split()[-1]) if style.split()[-1].isdigit() else 1
                    doc.paragraphs[-1].clear()
                    doc.paragraphs.pop()
                    doc.add_heading(text, level=level)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_import_with_heading_chapters(client: TestClient) -> None:
    pad = "Prose continues here. " * 12
    paragraphs = [
        ("Heading 1", "Chapter 1: The Beginning"),
        ("", "First paragraph. " + pad),
        ("", "Second paragraph. " + pad),
        ("Heading 1", "Chapter 2: Onwards"),
        ("", "Body of chapter two. " + pad),
    ]
    data = _build_docx_bytes(paragraphs)
    r = client.post(
        "/api/translate/upload",
        data={"title": "DOCX Novel"},
        files={
            "file": (
                "n.docx",
                BytesIO(data),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert r.status_code == 200, r.text
    body = r.json()
    assert body["source_type"] == "docx"
    assert body["detected_encoding"] == "docx"
    novel_id = body["novel_id"]

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = list(
        conn.execute(
            "SELECT chapter_num, title_zh, original_text FROM chapters "
            "WHERE novel_id = ? ORDER BY chapter_num",
            (novel_id,),
        )
    )
    conn.close()
    assert len(rows) == 2, rows
    assert "Chapter 1" in (rows[0]["title_zh"] or "")
    assert "Chapter 2" in (rows[1]["title_zh"] or "")
    assert "First paragraph" in rows[0]["original_text"]
    assert "Body of chapter two" in rows[1]["original_text"]


def test_docx_import_cjk_punctuation_preserved(client: TestClient) -> None:
    """Reads-as-Chinese-content DOCX — verify CJK chars + 第N章 split works
    end-to-end. python-docx is utf-8 internally so this is mostly a regression
    test for our own paragraph-joining."""
    pad = "情节继续展开，主角站在山顶上眺望远方。" * 20
    paragraphs = [
        ("Heading 1", "第一章 序幕"),
        ("", "第一段内容。" + pad),
        ("Heading 1", "第二章 行动"),
        ("", "第二章正文。" + pad),
    ]
    data = _build_docx_bytes(paragraphs)
    r = client.post(
        "/api/translate/upload",
        data={"title": "CJK"},
        files={"file": ("c.docx", BytesIO(data), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200, r.text

    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    rows = list(
        conn.execute(
            "SELECT chapter_num, title_zh, original_text FROM chapters "
            "WHERE novel_id = ? ORDER BY chapter_num",
            (r.json()["novel_id"],),
        )
    )
    conn.close()
    assert len(rows) == 2, rows
    assert "第一章" in (rows[0]["title_zh"] or "")
    assert "山顶" in rows[0]["original_text"]


def test_docx_import_rejects_corrupt_archive(client: TestClient) -> None:
    r = client.post(
        "/api/translate/upload",
        data={"title": "Bad"},
        files={"file": ("broken.docx", BytesIO(b"not a real docx"), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 400, r.text
    assert "could not parse .docx" in r.json()["detail"]


def test_bulk_endpoint_rejects_docx(client: TestClient) -> None:
    data = _build_docx_bytes([("", "body" * 100)])
    r = client.post(
        "/api/translate/bulk",
        data={"title": "Y"},
        files=[("files", ("a.docx", BytesIO(data), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"))],
    )
    assert r.status_code == 400, r.text
